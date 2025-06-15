"""
文件处理工具 - 支持多种格式的文本提取
"""
import os
import mimetypes
from typing import Optional, Tuple
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# 可选依赖的导入
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False


class FileProcessor:
    """文件处理器类"""
    
    SUPPORTED_EXTENSIONS = ['.txt', '.md', '.pdf', '.docx']
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    @classmethod
    def is_supported_file(cls, filename: str) -> bool:
        """检查文件是否支持"""
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def get_file_type(cls, filename: str) -> str:
        """获取文件类型"""
        ext = Path(filename).suffix.lower()
        return ext.lstrip('.')
    
    @classmethod
    def validate_file_size(cls, file_path: str) -> bool:
        """验证文件大小"""
        try:
            size = os.path.getsize(file_path)
            return size <= cls.MAX_FILE_SIZE
        except OSError:
            return False
    
    @classmethod
    def extract_text(cls, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        从文件中提取文本内容
        
        Returns:
            Tuple[bool, str, Optional[str]]: (是否成功, 提取的文本内容, 错误信息)
        """
        if not os.path.exists(file_path):
            return False, "", "文件不存在"
        
        if not cls.validate_file_size(file_path):
            return False, "", f"文件大小超过限制 ({cls.MAX_FILE_SIZE / 1024 / 1024:.1f}MB)"
        
        filename = os.path.basename(file_path)
        if not cls.is_supported_file(filename):
            return False, "", f"不支持的文件类型: {Path(filename).suffix}"
        
        try:
            file_type = cls.get_file_type(filename)
            
            if file_type in ['txt', 'md']:
                return cls._extract_text_file(file_path)
            elif file_type == 'pdf':
                return cls._extract_pdf(file_path)
            elif file_type == 'docx':
                return cls._extract_docx(file_path)
            else:
                return False, "", f"暂不支持的文件类型: {file_type}"
                
        except Exception as e:
            logger.error(f"处理文件 {file_path} 时出错: {str(e)}")
            return False, "", f"文件处理失败: {str(e)}"
    
    @classmethod
    def _extract_text_file(cls, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """提取文本文件内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return True, content, None
                except UnicodeDecodeError:
                    continue
            
            return False, "", "无法解码文件，请检查文件编码"
            
        except Exception as e:
            return False, "", f"读取文本文件失败: {str(e)}"
    
    @classmethod
    def _extract_pdf(cls, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """提取PDF文件内容"""
        if not PDF_AVAILABLE:
            return False, "", "PDF处理库未安装，请安装 PyPDF2: pip install PyPDF2"
        
        try:
            content = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
            
            if not content.strip():
                return False, "", "PDF文件中未找到可提取的文本内容"
            
            return True, content.strip(), None
            
        except Exception as e:
            return False, "", f"PDF文件处理失败: {str(e)}"
    
    @classmethod
    def _extract_docx(cls, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """提取DOCX文件内容"""
        if not DOCX_AVAILABLE:
            return False, "", "DOCX处理库未安装，请安装 python-docx: pip install python-docx"
        
        try:
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            if not content.strip():
                return False, "", "DOCX文件中未找到可提取的文本内容"
            
            return True, content.strip(), None
            
        except Exception as e:
            return False, "", f"DOCX文件处理失败: {str(e)}"


def extract_file_content(file_path: str) -> Tuple[bool, str, Optional[str]]:
    """
    文件内容提取的便捷函数
    
    Args:
        file_path: 文件路径
        
    Returns:
        Tuple[bool, str, Optional[str]]: (是否成功, 提取的文本, 错误信息)
    """
    return FileProcessor.extract_text(file_path)


def get_supported_file_types() -> list:
    """获取支持的文件类型列表"""
    return [ext.lstrip('.') for ext in FileProcessor.SUPPORTED_EXTENSIONS]


def is_file_supported(filename: str) -> bool:
    """检查文件是否支持"""
    return FileProcessor.is_supported_file(filename) 